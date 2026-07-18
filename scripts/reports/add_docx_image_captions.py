from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


CAPTIONS = [
    "图1：H3 网格与高德逆地理编码地名增强示意图",
    "图2：异常活动页面与 AI 分析结果截图",
    "图3：AI 辅助异常解释结果局部截图",
    "图4：北京夜间活力地图主界面截图",
    "图5：结果可信度与短时预测页面截图",
    "图6：Top 网格小时变化与周末效应图",
]


def _latest_report() -> Path:
    reports = Path("reports")
    candidates = [
        path
        for path in reports.glob("*.docx")
        if "_before_captions" not in path.stem and not path.stem.endswith("_captions")
    ]
    candidates = sorted(candidates, key=lambda path: path.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError("reports 目录下没有找到 docx 文件。")
    return candidates[0]


def _paragraph_after(paragraph):
    new_paragraph = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph)
    return paragraph.__class__(new_paragraph, paragraph._parent)


def _image_runs(paragraph):
    return [run for run in paragraph.runs if run._element.xpath(".//pic:pic")]


def _split_multi_image_paragraph(paragraph) -> None:
    image_runs = _image_runs(paragraph)
    if len(image_runs) <= 1:
        return

    current_paragraph = paragraph
    for run in image_runs[1:]:
        new_element = OxmlElement("w:p")
        paragraph_properties = paragraph._p.pPr
        if paragraph_properties is not None:
            new_element.append(deepcopy(paragraph_properties))
        current_paragraph._p.addnext(new_element)
        new_paragraph = paragraph.__class__(new_element, paragraph._parent)
        new_paragraph._p.append(run._element)
        current_paragraph = new_paragraph


def _is_picture_paragraph(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//pic:pic"))


def _apply_caption_format(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "宋体")
    rfonts.set(qn("w:hAnsi"), "宋体")


def _next_caption_paragraph(paragraph):
    next_element = paragraph._p.getnext()
    if next_element is None or next_element.tag != qn("w:p"):
        return None
    next_paragraph = paragraph.__class__(next_element, paragraph._parent)
    if next_paragraph.text.strip().startswith("图"):
        return next_paragraph
    return None


def main() -> None:
    report_path = _latest_report()
    backup_path = report_path.with_name(f"{report_path.stem}_before_captions.docx")
    if not backup_path.exists():
        shutil.copy2(report_path, backup_path)

    doc = Document(str(report_path))
    for paragraph in list(doc.paragraphs):
        _split_multi_image_paragraph(paragraph)

    picture_paragraphs = [paragraph for paragraph in doc.paragraphs if _is_picture_paragraph(paragraph)]
    if len(picture_paragraphs) != len(CAPTIONS):
        raise RuntimeError(f"检测到 {len(picture_paragraphs)} 张图片，但配置了 {len(CAPTIONS)} 条图注。")

    for picture_paragraph, caption in reversed(list(zip(picture_paragraphs, CAPTIONS))):
        caption_paragraph = _next_caption_paragraph(picture_paragraph) or _paragraph_after(picture_paragraph)
        _apply_caption_format(caption_paragraph, caption)

    try:
        doc.save(str(report_path))
        print(f"已保存：{report_path}")
    except PermissionError:
        fallback = report_path.with_name(f"{report_path.stem}_{datetime.now():%H%M%S}_captions.docx")
        doc.save(str(fallback))
        print(f"原文件可能被 Word 占用，已另存为：{fallback}")
    print(f"备份文件：{backup_path}")
    print(f"已处理图注：{len(CAPTIONS)} 条")


if __name__ == "__main__":
    main()
