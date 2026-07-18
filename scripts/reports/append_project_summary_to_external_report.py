from __future__ import annotations

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REPORT_DIR = Path(r"D:\uns_2\report")

SECTION_ITEMS = [
    ("h1", "十三、项目总结与自评"),
    ("h2", "13.1 项目文字总结"),
    (
        "p",
        "本项目最终形成了一个面向北京城市夜间活动观察的出租车轨迹分析系统。系统以北京 T-Drive 出租车 GPS "
        "轨迹为基础，将连续轨迹点整理为轨迹活动段，并进一步聚合到 H3 活动网格中，构建夜间活动规模、时间结构、"
        "流动特征、POI 功能画像、异常活动和可信度分析等结果。前端通过高德地图展示北京夜间出租车活动的空间分布，"
        "使抽象的数据处理结果能够以网格、地名、颜色等级和详情卡片的方式被直观理解。",
    ),
    (
        "p",
        "从项目逻辑看，当前版本不再把出租车数据简单解释为夜生活消费强度，而是将其定位为“夜间出租车活动与城市活力监测”。"
        "这一定位更符合 T-Drive 数据本身的字段条件：数据能够真实记录车辆轨迹和时间变化，但不包含完整乘客订单、费用和目的信息。"
        "因此，报告和系统统一使用“活动段”“活动起点”“活动终点”“活动网格”等概念，使分析边界更加清楚，也避免了对数据含义的过度解释。",
    ),
    (
        "p",
        "项目最终实现了从数据转换、空间聚合、活力评分、异常识别、POI 增强、结果解释到前端可视化的完整流程。地图、"
        "异常分析、网格画像和可信度页面并不是独立堆叠的功能，而是共同服务于一个目标：帮助使用者识别北京夜间出租车活动在哪里集中、"
        "何时集中、可能与哪些城市功能有关，以及哪些结果需要谨慎解释。",
    ),
    ("h2", "13.2 项目自评"),
    (
        "p",
        "从完成度上看，项目已经具备较完整的数据分析与展示闭环。其优点主要体现在三个方面。第一，项目能够根据数据特征调整研究对象，"
        "没有机械套用原有订单数据叙事，而是围绕北京 T-Drive 轨迹数据重新建立了活动段和活动网格体系。第二，系统功能与研究目的之间的关系较为明确，"
        "H3 网格负责空间表达，高德逆地理编码负责地名可读性，POI 画像负责功能解释，异常检测负责发现局部波动，可信度分析负责说明结果边界。"
        "第三，前端展示较好地降低了结果理解门槛，使用连续颜色、核心网格标记、点击高亮和详情卡片，使数据结果更接近一个可交互的城市监测工具。",
    ),
    (
        "p",
        "项目也存在一些不足。首先，T-Drive 数据发布时间较早，时间跨度和车辆样本有限，不能直接代表当前北京全部夜间出行结构。"
        "其次，活动段由规则推断得到，并不等同于真实乘客上下车订单，因此无法分析费用、真实目的地和乘客行为。再次，POI 信息主要来自高德逆地理编码返回的网格中心附近信息，"
        "仍属于近似增强，而不是完整 POI 普查。最后，异常检测目前主要依据数据内部历史基线，能够提示“异常波动”，但还不能独立确认真实事件原因。"
        "总体来看，本项目更适合作为夜间出租车活动监测和城市空间观察的辅助系统，而不是对夜间经济或乘客出行行为的完整测量。",
    ),
    ("h2", "13.3 项目完成过程中遇到的困难及解决方法"),
    (
        "p",
        "项目过程中最核心的困难是数据来源与项目定位之间的矛盾。项目早期使用的是纽约出租车数据，原因是这类开放数据较容易获取，"
        "字段完整，包含上下车区域、时间、距离和费用等信息，适合快速搭建夜间出行分析原型。但随着项目希望将应用地改回中国，"
        "原有纽约案例就暴露出明显问题：一方面，纽约数据难以支撑中国城市应用叙事；另一方面，国内公开、可直接使用、字段完整的打车订单数据并不好获取，"
        "许多数据要么缺少空间边界，要么缺少订单状态，要么无法公开下载。",
    ),
    (
        "p",
        "为解决这一问题，项目转向北京 T-Drive 出租车 GPS 轨迹数据。这个选择并不是简单替换城市名称，而是带来了研究对象的重新定义。"
        "T-Drive 提供的是车辆轨迹点，而不是订单表，因此项目不能继续沿用“上车量、下车量、订单热区”的完整订单叙事。"
        "解决方法是将轨迹点按车辆、时间连续性、速度变化和停留状态切分为轨迹活动段，把活动段首尾位置作为夜间出租车活动的近似起点和终点，"
        "再用 H3 网格进行空间聚合。这样既保留了出租车活动的时空信息，也避免了把轨迹数据误读成真实订单数据。",
    ),
    (
        "p",
        "数据切换之后，项目定位也随之调整。原先的“夜生活分析”容易让人认为系统直接衡量消费、娱乐或人群活动，但 T-Drive 数据更适合反映出租车运行视角下的夜间活动强度。"
        "因此，项目最终明确为“北京夜间出租车活动与城市活力监测系统”，目标用户也更偏向城市交通管理、夜间运行监测、空间治理和城市研究场景。"
        "这个定位使后续功能有了统一解释：异常分析用于发现短时需求波动，POI 画像用于辅助解释网格功能，可信度分析用于说明模型结果是否稳定，"
        "地图交互用于支持空间观察和案例展示。",
    ),
    (
        "p",
        "第二个困难是空间表达和地名可读性问题。H3 网格虽然适合统一空间分析，但网格编号本身难以被读者理解；如果网格分辨率过高，"
        "前端还会出现网格数量过多、渲染卡顿和底图被覆盖的问题。解决方法是降低 H3 分辨率到适合展示和分析的层级，并通过高德逆地理编码为网格中心点补充区县、"
        "街道、商圈、附近 POI 和 AOI 信息。这样既控制了地图性能，又让结果从抽象编号转化为“朝阳区 · 大屯街道”等更容易理解的空间名称。",
    ),
    (
        "p",
        "第三个困难是系统功能较多但容易分散。项目中包含活力评分、异常检测、POI 分析、预测、可信度和 AI 辅助解释等功能，如果缺少明确主线，"
        "就会变成多个工具的简单堆叠。解决方法是在报告和前端中统一强调“夜间出租车活动监测”这一核心目标，并让每个功能都对应一个解释任务："
        "评分回答哪里更活跃，小时曲线回答何时更活跃，POI 回答可能与什么城市功能有关，异常检测回答哪里出现突增，可信度回答结果是否稳定。"
        "经过这一调整，项目整体逻辑更加顺畅，也更适合作为完整论文报告进行呈现。",
    ),
]


def latest_report() -> Path:
    candidates = sorted(REPORT_DIR.glob("*.docx"), key=lambda path: path.stat().st_mtime, reverse=True)
    candidates = [path for path in candidates if "before_project_summary" not in path.stem and "已追加总结" not in path.stem]
    if not candidates:
        raise FileNotFoundError(f"No editable docx found in {REPORT_DIR}")
    return candidates[0]


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def apply_paragraph_font(paragraph, size: float, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size, bold=bold)


def insert_before(anchor, text: str, style_name: str):
    new_element = OxmlElement("w:p")
    anchor._p.addprevious(new_element)
    paragraph = anchor.__class__(new_element, anchor._parent)
    paragraph.style = style_name
    paragraph.add_run(text)
    return paragraph


def apply_body_format(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)


def main() -> None:
    report_path = latest_report()
    backup_path = report_path.with_name(f"{report_path.stem}_before_project_summary.docx")
    if not backup_path.exists():
        shutil.copy2(report_path, backup_path)

    doc = Document(str(report_path))

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "十、可信度分析与局限性":
            paragraph.clear()
            paragraph.style = "Heading 1"
            paragraph.add_run("十一、可信度分析与局限性")
            apply_paragraph_font(paragraph, 14, bold=True)
        elif text == "一、结论与展望":
            paragraph.clear()
            paragraph.style = "Heading 1"
            paragraph.add_run("十二、结论与展望")
            apply_paragraph_font(paragraph, 14, bold=True)

    anchor = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "附录"), None)
    if anchor is None:
        raise RuntimeError("未找到“附录”段落，无法确定插入位置。")

    # If this script has been run before, remove the existing inserted section
    # between "12.3 后续工作" content and the appendix before writing it again.
    paragraphs = list(doc.paragraphs)
    anchor_index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is anchor._p)
    remove_after_index = None
    for index, paragraph in enumerate(paragraphs[:anchor_index]):
        if paragraph.text.strip().startswith("后续可继续引入更多月份或更新年份"):
            remove_after_index = index
    if remove_after_index is not None:
        for paragraph in paragraphs[remove_after_index + 1 : anchor_index]:
            paragraph._element.getparent().remove(paragraph._element)

    anchor = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "附录"), None)
    if anchor is None:
        raise RuntimeError("未找到“附录”段落，无法确定插入位置。")

    for kind, text in SECTION_ITEMS:
        if kind == "h1":
            paragraph = insert_before(anchor, text, "Heading 1")
            apply_paragraph_font(paragraph, 14, bold=True)
        elif kind == "h2":
            paragraph = insert_before(anchor, text, "Heading 2")
            apply_paragraph_font(paragraph, 12, bold=True)
        else:
            paragraph = insert_before(anchor, text, "Normal")
            apply_body_format(paragraph)
            apply_paragraph_font(paragraph, 12, bold=False)

    try:
        doc.save(str(report_path))
        print(f"saved: {report_path}")
    except PermissionError:
        fallback = report_path.with_name(f"{report_path.stem}_{datetime.now():%H%M%S}_已追加总结.docx")
        doc.save(str(fallback))
        print(f"locked_saved_as: {fallback}")
    print(f"backup: {backup_path}")


if __name__ == "__main__":
    main()
