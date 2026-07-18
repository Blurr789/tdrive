from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


OUTPUT_PATH = Path(r"D:\uns_2\report\北京T-Drive项目讲解视频文稿与功能演示顺序.docx")


SECTIONS = [
    (
        "一、开场与项目定位",
        "打开前端页面，停留在“活动地图”首页，先不要操作地图。",
        "大家好，我这次展示的项目是“北京 T-Drive 夜间出租车活动与城市活力监测系统”。本项目使用北京 T-Drive 出租车 GPS 轨迹数据，围绕北京夜间出租车活动的空间分布、时间变化、异常波动和城市功能解释进行分析。需要先说明的是，T-Drive 并不是完整订单数据，它没有真实乘客上车、下车、费用和目的地字段。因此本项目没有直接把数据解释为出租车订单，而是根据车辆轨迹点的时间连续性、移动距离和停留状态，构造出轨迹活动段，再把活动段聚合到 H3 活动网格中，用于观察北京夜间出租车活动和城市活力特征。",
    ),
    (
        "二、总体指标展示",
        "在活动地图页面左侧停留，指向活动网格数量、异常活动数量和最高分等总体指标。",
        "首先看左侧总体统计。这里展示的是当前参与分析的活动网格数量、系统识别出的异常活动数量，以及最高夜间活力分数。这些指标可以帮助我们快速了解本次数据处理后的总体规模。当前系统不是按行政区整体分析，而是使用 H3 活动网格，因此能够比区县尺度更细致地观察北京夜间出租车活动的空间差异。",
    ),
    (
        "三、北京夜间活力地图",
        "观察右侧高德地图和 H3 网格颜色，说明颜色渐变含义；可以适当缩放或平移地图。",
        "右侧是高德地图底图，上面叠加的是 H3 活动网格。颜色使用连续暖色渐变表达活动强度，颜色越深、越偏橙红，说明该网格夜间出租车活动越强。相比使用很多完全不同的颜色，渐变色更适合表达同一指标的强弱变化，也更方便读者从地图上判断热点区域的空间分布。",
    ),
    (
        "四、核心网格标记",
        "指向地图上的 CORE 标记，说明为什么只显示 Top 10 核心网格。",
        "地图上只显示 Top 10 的 CORE 核心网格标记。这是为了避免所有网格标签同时显示造成信息过载。H3 网格数量较多，如果所有标签都显示出来，会覆盖底图，也会影响读图效率。因此系统只突出最核心的活动网格，让用户优先关注夜间出租车活动最集中的位置。",
    ),
    (
        "五、时间滑块与小时变化",
        "拖动夜间时间滑块，依次展示综合、20 点、21 点、23 点、0 点、2 点等状态。",
        "接下来切换时间滑块。项目把夜间时段定义为 20 点到次日 3 点。综合视图可以观察整夜总体空间格局，小时视图则可以观察不同时间段的活动变化。比如晚间时段可能更体现商务、餐饮和休闲活动，深夜和凌晨则可能更多反映返程、接驳或交通枢纽需求。通过时间滑块，系统不仅能告诉我们哪里活跃，也能展示这些区域在夜间不同时段的变化节奏。",
    ),
    (
        "六、点击网格与详情卡片",
        "点击一个颜色较深或带 CORE 标记的网格，展示区块高亮、弹窗和左侧详情卡片。",
        "现在点击一个核心网格。被选中的网格会在地图上高亮，弹窗和左侧详情卡片会显示该网格的信息，包括网格名称、活力分数、活动类型和相关统计。这里的名称不是原始 H3 编号，而是通过高德逆地理编码，把网格中心点转换成区县、街道、商圈或附近 POI。这样可以把抽象网格和真实北京空间对应起来，降低结果理解门槛。",
    ),
    (
        "七、搜索定位功能",
        "在搜索框输入“朝阳”“机场”“学院路”等关键词，展示搜索结果并点击定位。",
        "这里还可以使用搜索功能。例如输入“朝阳”“机场”或“学院路”，系统会在活动网格编号、展示名称、区县、街道等字段中匹配结果。这个功能主要解决 H3 网格难以阅读的问题。用户不需要记住复杂编号，而是可以通过熟悉的北京地名找到对应活动网格。",
    ),
    (
        "八、异常活动页面",
        "切换到“异常活动”页面，浏览异常卡片列表，点击其中一条异常记录。",
        "下面切换到异常活动页面。这里列出了系统识别出的夜间活动异常。异常检测使用同一网格、同一小时的历史活动基线，通过 median 和 MAD 方法计算异常强度。每条异常记录会展示异常发生的日期、小时、地点、活动量、历史基线和异常等级。这个页面的作用是帮助使用者发现短时需求突增或局部活动波动。",
    ),
    (
        "九、AI 辅助异常解释",
        "在异常活动页面点击一条异常后，展示 AI 分析区域；如有按钮，可点击“重新分析”。",
        "异常页面还提供 AI 辅助解释。需要强调的是，AI 分析严格限制在项目内部数据范围内，不直接编造外部新闻、天气或真实活动事件。它会根据异常活动量、历史基线、前后小时变化、周末效应、网格类型和 POI 特征，给出可能原因、数据证据和建议核查方向。因此它更像是辅助分析工具，而不是事实确认工具。",
    ),
    (
        "十、网格画像页面",
        "切换到“网格画像”页面，展示网格列表、筛选、搜索和某个网格详情。",
        "接下来进入网格画像页面。这个页面用于解释每个活动网格的功能特征。用户可以搜索、筛选和查看不同网格。网格画像中包括活力分数、活动类型、POI 多样性、夜间相关 POI、交通设施相关指标、活动起终点结构等内容。它回答的是：一个网格为什么可能活跃，它周边有哪些城市功能，它更像交通接驳区、商务活动区，还是混合夜间活动区。",
    ),
    (
        "十一、POI 与地名增强说明",
        "在网格画像或详情卡片中指向 POI 多样性、商圈、街道、附近 POI 等字段。",
        "这里的 POI 和地名信息来自高德逆地理编码。系统对每个 H3 网格计算中心点，然后获取区县、街道、商圈、附近 POI 和 AOI 信息。这样做有两个作用：第一，让网格有可读名称；第二，为活力结果提供功能解释。例如某些区域活动强，可能与交通设施、餐饮休闲、商务办公或居住返程有关。POI 画像不能等同于完整空间普查，但能够为理解网格功能提供辅助依据。",
    ),
    (
        "十二、可信度页面",
        "切换到“可信度”页面，展示评分方法对比、预测指标和下一小时预测卡片。",
        "然后进入可信度页面。这个页面主要说明结果是否稳定，以及模型结果应该如何谨慎理解。项目不仅使用人工权重计算 Night Vitality Score，也生成了熵权法和 PCA 的评分对比。不同方法排序不完全一致，说明夜间活力不是单一维度指标。因此系统采用人工权重作为主解释指标，同时用其他方法作为稳健性参考。",
    ),
    (
        "十三、短时预测结果",
        "在可信度页面展示 Random Forest、上一小时参考和下一小时预测结果。",
        "可信度页面还包含短时预测部分。系统使用历史小时活动量构建下一小时活动预测，用于判断夜间活动是否具有一定时间规律。这里的预测结果不是项目的主要结论，而是作为可信度辅助分析。它可以帮助我们理解哪些区域的短时活动变化比较稳定，哪些区域可能更容易受偶发因素影响。",
    ),
    (
        "十四、自动报告与结果输出",
        "如果前端有报告内容页面则展示；如果没有，可口头说明 outputs 和 Word 报告已生成。",
        "除前端页面外，项目还生成了多类结果文件，包括 H3 地图图层、网格评分、异常记录、POI 画像、评分方法对比、预测结果和自动报告。这说明项目不只是一个前端展示页面，而是形成了从数据处理、算法分析、结果输出到可视化展示的完整流程。",
    ),
    (
        "十五、项目总结与局限",
        "回到活动地图页面，停留在北京整体地图视图，进行总结收尾。",
        "最后总结一下。本项目的核心价值，是把北京 T-Drive 轨迹数据转化成一个可观察、可解释、可交互的夜间出租车活动监测系统。它解决了三个关键问题：第一，如何在没有完整订单数据的情况下构造合理的活动分析单元；第二，如何用 H3 网格和高德逆地理编码把轨迹数据映射到真实北京空间；第三，如何通过活力评分、异常检测、POI 画像、AI 辅助解释和可信度分析，让所有功能共同服务于夜间出租车活动监测这一主题。当然，项目也有局限。T-Drive 数据时间较早，活动段不等同于真实乘客订单，POI 信息也只是网格中心附近的辅助解释。因此本系统更适合作为城市夜间交通和空间活力观察的辅助工具，而不是完整乘客出行行为研究。以上就是本项目的主要展示。",
    ),
]


def set_font(run, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color or RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_paragraph(doc: Document, text: str, size: float = 12, bold: bool = False, first_indent: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def add_label_text(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(label)
    set_font(label_run, size=11, bold=True, color=RGBColor(31, 78, 121))
    text_run = paragraph.add_run(text)
    set_font(text_run, size=11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("北京 T-Drive 夜间出租车活动与城市活力监测系统\n项目讲解视频文稿与功能演示顺序")
    set_font(title_run, size=18, bold=True)

    add_paragraph(doc, "建议视频时长：6-8 分钟。录制方式建议采用“先说当前页面作用，再操作，再解释结果”的节奏。", size=11)

    add_heading(doc, "一、整体演示路线", 1)
    route = "首页与项目定位 → 总体指标 → 活力地图 → 核心网格 → 时间滑块 → 网格详情 → 搜索定位 → 异常活动 → AI 辅助解释 → 网格画像 → POI 与地名增强 → 可信度分析 → 短时预测 → 输出与报告 → 总结局限"
    add_paragraph(doc, route, size=11, first_indent=True)

    add_heading(doc, "二、分段讲解文稿", 1)
    for title_text, action, script in SECTIONS:
        add_heading(doc, title_text, 2)
        add_label_text(doc, "演示操作：", action)
        add_label_text(doc, "讲解文稿：", script)

    add_heading(doc, "三、录制提醒", 1)
    tips = [
        "录制前先启动后端和前端，并确认高德地图能够加载。",
        "讲解地图时不要频繁缩放，重点展示颜色、核心网格、时间滑块和点击详情。",
        "异常活动页面至少点开一条异常，并展示 AI 分析结果。",
        "网格画像页面重点展示 POI 多样性、地名增强和网格功能解释。",
        "可信度页面重点说明评分方法对比和短时预测只是辅助判断，不要过度解释。",
    ]
    for tip in tips:
        add_paragraph(doc, f"• {tip}", size=11)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
