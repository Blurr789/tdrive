from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "outputs" / "beijing_tdrive"
REPORTS = ROOT / "reports"
REPORT_PATH = REPORTS / "北京T-Drive夜间出租车活动与城市活力监测系统研究报告.docx"

CN_NUMBERS = "一二三四五六七八九十十一十二十三十四"

REGION_TYPE_LABELS = {
    "nightlife_core": "夜间活动核心区",
    "transport_hub_like": "交通枢纽型网格",
    "return_destination": "夜间到达热点",
    "departure_hotspot": "夜间出发热点",
    "mixed_evening_area": "混合夜间活动区",
    "low_activity": "低活跃区",
    "no_activity": "暂无活跃记录",
    "insufficient_data": "数据不足",
}


def set_run_font(run, size: int = 12, bold: bool | None = None, font_cn: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_paragraph(
    doc: Document,
    text: str = "",
    size: int = 12,
    bold: bool = False,
    align=None,
    spacing_after: int = 6,
    first_line_indent: bool = True,
    font_cn: str = "宋体",
):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    if text and first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, font_cn=font_cn)
    return paragraph


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        add_paragraph(doc, first_line_indent=False)
    add_paragraph(
        doc,
        "基于北京 T-Drive 轨迹数据的夜间出租车活动与城市活力监测系统研究",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        spacing_after=20,
        first_line_indent=False,
        font_cn="黑体",
    )
    for _ in range(5):
        add_paragraph(doc, first_line_indent=False)
    add_paragraph(doc, "Urban NightSense 项目完整报告", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    add_paragraph(doc, "2023111748 曹黎悦", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    doc.add_page_break()


def set_heading_styles(doc: Document) -> None:
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_toc_page(doc: Document) -> None:
    add_paragraph(doc, "目录", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12, first_line_indent=False)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(field_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'
    instr_run._r.append(instr_text)

    separate_run = paragraph.add_run()
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(field_separate)

    placeholder = paragraph.add_run("请在 Word 中右键选择“更新域/更新整个目录”，生成目录页码。")
    set_run_font(placeholder, size=11)

    end_run = paragraph.add_run()
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(field_end)

    update = doc.settings.element.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        doc.settings.element.append(update)
    update.set(qn("w:val"), "true")
    doc.add_page_break()


def add_section(doc: Document, index: int, title: str) -> None:
    paragraph = add_paragraph(doc, f"{CN_NUMBERS[index - 1]}、{title}", size=14, bold=True, spacing_after=8, first_line_indent=False)
    paragraph.style = doc.styles["Heading 1"]


def add_abstract_heading(doc: Document) -> None:
    paragraph = add_paragraph(doc, "摘要", size=14, bold=True, spacing_after=8, first_line_indent=False)
    paragraph.style = doc.styles["Heading 1"]


def add_subsection(doc: Document, number: str, title: str) -> None:
    paragraph = add_paragraph(doc, f"{number} {title}", size=12, bold=True, spacing_after=6, first_line_indent=False)
    paragraph.style = doc.styles["Heading 2"]


def add_appendix_heading(doc: Document) -> None:
    paragraph = add_paragraph(doc, "附录", size=14, bold=True, spacing_after=8, first_line_indent=False)
    paragraph.style = doc.styles["Heading 1"]


def force_all_text_black(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
    for style in doc.styles:
        if hasattr(style, "font"):
            try:
                style.font.color.rgb = RGBColor(0, 0, 0)
            except (AttributeError, TypeError):
                pass


def add_caption(doc: Document, text: str) -> None:
    add_paragraph(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=6, first_line_indent=False)


def add_image_placeholder(doc: Document, label: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(label)
    set_run_font(run, size=10, bold=True)
    add_paragraph(doc, first_line_indent=False, spacing_after=8)


def add_table(doc: Document, rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=10, bold=(i == 0))
    add_paragraph(doc, first_line_indent=False)


def fmt_int(value: Any) -> str:
    return f"{int(float(value or 0)):,}"


def fmt_float(value: Any, digits: int = 2) -> str:
    return f"{float(value or 0):.{digits}f}"


def fmt_pct(value: Any, digits: int = 1) -> str:
    return f"{float(value or 0) * 100:.{digits}f}%"


def clean_text(value: Any) -> str:
    if pd.isna(value):
        return ""
    return str(value)


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def load_name_lookup() -> pd.DataFrame:
    path = OUTPUT / "grid_name_lookup.csv"
    if not path.exists():
        return pd.DataFrame(columns=["spatial_unit", "display_name", "district", "township"])
    lookup = pd.read_csv(path)
    lookup["spatial_unit"] = lookup["spatial_unit"].astype(str)
    return lookup


def load_scores() -> pd.DataFrame:
    scores = pd.read_csv(OUTPUT / "region_scores_phase2.csv")
    scores["spatial_unit"] = scores["spatial_unit"].astype(str)
    lookup = load_name_lookup()
    if not lookup.empty:
        keep = [
            "spatial_unit",
            "display_name",
            "district",
            "township",
            "business_area",
            "nearest_poi",
            "nearest_aoi",
            "name_source",
        ]
        scores = scores.merge(lookup[[col for col in keep if col in lookup.columns]], on="spatial_unit", how="left")
    scores["display_name"] = scores["display_name"].fillna(scores.get("zone", scores["spatial_unit"]))
    scores["district"] = scores["district"].fillna(scores.get("borough", "北京市"))
    return scores.sort_values("night_vitality_score", ascending=False).reset_index(drop=True)


def load_anomalies(scores: pd.DataFrame) -> pd.DataFrame:
    path = OUTPUT / "anomalies_attributed.csv"
    if not path.exists():
        return pd.DataFrame()
    anomalies = pd.read_csv(path)
    anomalies["spatial_unit"] = anomalies["spatial_unit"].astype(str)
    names = scores[["spatial_unit", "display_name", "district"]].drop_duplicates("spatial_unit")
    return anomalies.merge(names, on="spatial_unit", how="left")


def top_region_rows(scores: pd.DataFrame, limit: int = 10) -> list[list[Any]]:
    rows = [["排名", "活动网格名称", "行政区", "活力分数", "POI 多样性", "区域类型"]]
    for index, row in enumerate(scores.head(limit).to_dict("records"), start=1):
        rows.append(
            [
                index,
                clean_text(row.get("display_name")),
                clean_text(row.get("district")),
                fmt_float(row.get("night_vitality_score"), 2),
                fmt_pct(row.get("poi_diversity")),
                REGION_TYPE_LABELS.get(clean_text(row.get("region_type")), clean_text(row.get("region_type"))),
            ]
        )
    return rows


def typical_region_rows(scores: pd.DataFrame) -> list[list[Any]]:
    candidates = []
    labels = [
        ("空港与交通活动型", scores[scores["display_name"].str.contains("机场|空港", na=False)].head(1)),
        ("商务活动型", scores[scores["display_name"].str.contains("建外|高碑店|大屯", na=False)].head(1)),
        ("高校与科研活动型", scores[scores["display_name"].str.contains("学院路|万寿路", na=False)].head(1)),
        ("城市核心混合型", scores[scores["display_name"].str.contains("和平里|麦子店|呼家楼", na=False)].head(1)),
    ]
    for label, table in labels:
        if not table.empty:
            row = table.iloc[0]
            candidates.append(
                [
                    label,
                    row["display_name"],
                    fmt_float(row["night_vitality_score"], 2),
                    fmt_pct(row.get("poi_diversity")),
                    fmt_int(row.get("total_activity")),
                    REGION_TYPE_LABELS.get(str(row.get("region_type")), str(row.get("region_type"))),
                ]
            )
    return [["类型", "代表网格", "活力分数", "POI 多样性", "活动总量", "识别类型"], *candidates]


def anomaly_rows(anomalies: pd.DataFrame, limit: int = 8) -> list[list[Any]]:
    rows = [["网格名称", "日期", "小时", "活动量", "异常强度", "可能原因"]]
    if anomalies.empty:
        return rows
    for row in anomalies.head(limit).to_dict("records"):
        rows.append(
            [
                clean_text(row.get("display_name") or row.get("spatial_unit")),
                row.get("night_date"),
                f"{int(row.get('hour', 0))}:00",
                fmt_int(row.get("activity_count")),
                fmt_float(row.get("z_score"), 2),
                clean_text(row.get("possible_reason")),
            ]
        )
    return rows


def forecast_rows() -> list[list[Any]]:
    metrics_path = OUTPUT / "forecast_metrics.csv"
    rows = [["模型", "MAE", "RMSE", "MAPE"]]
    if not metrics_path.exists():
        return rows
    labels = {"previous_hour_baseline": "上一小时基线", "random_forest": "Random Forest"}
    for row in pd.read_csv(metrics_path).to_dict("records"):
        rows.append(
            [
                labels.get(row["model"], row["model"]),
                fmt_float(row["mae"], 2),
                fmt_float(row["rmse"], 2),
                fmt_pct(float(row["mape"]) / 100 if float(row["mape"]) > 1 else row["mape"]),
            ]
        )
    return rows


def score_method_rows(scores: pd.DataFrame) -> list[list[Any]]:
    path = OUTPUT / "score_method_comparison.csv"
    rows = [["活动网格", "人工权重", "熵权法", "PCA", "分数差异"]]
    if not path.exists():
        return rows
    table = pd.read_csv(path)
    table["spatial_unit"] = table["spatial_unit"].astype(str)
    names = scores[["spatial_unit", "display_name"]].drop_duplicates("spatial_unit")
    table = table.merge(names, on="spatial_unit", how="left")
    for row in table.head(8).to_dict("records"):
        rows.append(
            [
                clean_text(row.get("display_name") or row.get("spatial_unit")),
                fmt_float(row.get("manual_score"), 2),
                fmt_float(row.get("entropy_score"), 2),
                fmt_float(row.get("pca_score"), 2),
                fmt_float(row.get("score_spread"), 2),
            ]
        )
    return rows


def section_one(doc: Document) -> None:
    add_abstract_heading(doc)
    add_paragraph(
        doc,
        "本研究基于北京 T-Drive 出租车 GPS 轨迹数据，构建了一个面向夜间出租车活动识别、城市活力监测与异常活动分析的可视化系统。由于 T-Drive 数据并非完整订单数据，本文首先将连续轨迹点转换为可分析的轨迹活动段，并以活动段的起点与终点近似表达夜间出行活动的空间分布。随后，研究采用 H3 网格作为基础空间单元，结合高德逆地理编码为网格生成区县、街道、商圈和 POI 语义名称，在此基础上提取活动规模、时间结构、空间流动、POI 功能和交通枢纽等特征，构建 Night Vitality Score 评价指标，并识别夜间活动核心区、混合夜间活动区、交通枢纽型网格等区域类型。",
    )
    add_paragraph(
        doc,
        "系统层面，项目采用 Python 数据处理、Flask 后端 API、Vue 前端和高德地图 JS API 的三层架构，实现了夜间活力地图、异常活动分析、活动网格画像、可信度分析和 AI 辅助解释等功能。结果表明，北京夜间出租车活动在朝阳、海淀、顺义、东城、丰台等区域形成了若干核心活动网格，其中空港、商务区、高校与居住交通混合区呈现出不同的夜间活动机制。研究同时指出，T-Drive 数据时间跨度较短且缺少真实订单标签，相关结论应被理解为基于轨迹活动段的夜间出租车活动监测结果，而非完整乘客订单行为解释。",
    )
    add_paragraph(doc, "关键词：T-Drive；出租车轨迹；夜间活力；H3 网格；高德地图；异常检测；城市计算", first_line_indent=False)


def build_report() -> Path:
    REPORTS.mkdir(parents=True, exist_ok=True)
    summary = read_json(OUTPUT / "pipeline_summary.json")
    scores = load_scores()
    anomalies = load_anomalies(scores)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    set_heading_styles(doc)

    add_title_page(doc)
    add_toc_page(doc)
    section_one(doc)

    add_section(doc, 1, "引言")
    add_subsection(doc, "1.1", "研究背景")
    add_paragraph(doc, "夜间活动是城市运行的重要组成部分，能够反映交通供给、消费活动、居住返程、公共安全和空间功能结构等多方面信息。出租车在夜间具有较强的机动性和服务连续性，因此出租车轨迹数据可以成为观察城市夜间活动的重要数据来源。相比白天通勤活动，夜间出行更容易受到商业消费、交通枢纽、娱乐活动、居住返程和局部事件影响，因而也更需要精细化的空间监测与解释。")
    add_subsection(doc, "1.2", "研究问题")
    add_paragraph(doc, "本文关注四个问题：北京夜间出租车活动主要集中在哪些空间；不同网格的活力水平和功能类型如何区分；哪些时段和网格存在异常夜间活动；如何将抽象 H3 网格转换为可解释的北京街道、商圈与 POI 语义空间。")
    add_subsection(doc, "1.3", "研究贡献")
    add_paragraph(doc, "研究贡献主要包括：构建 T-Drive 轨迹活动段提取流程；采用 H3 网格实现 GPS 轨迹空间聚合；通过高德逆地理编码完成网格地名增强；基于高德 POI 和 AOI 信息构建功能画像；形成夜间活力评分、交通枢纽修正、异常检测和 AI 辅助解释的一体化系统；并最终实现可交互的北京夜间出租车活动监测前端。")
    add_subsection(doc, "1.4", "目标受众")
    add_paragraph(doc, "本项目的目标受众主要包括城市交通管理者、夜间经济与城市活力研究者、城市规划与空间治理人员，以及需要理解夜间出行活动分布的数据分析人员。系统并不面向普通公众提供消费推荐，也不试图替代真实交通调查，而是作为一种基于轨迹数据的辅助监测工具，帮助使用者快速识别夜间出租车活动集聚区、交通接驳热点、异常活动网格和需要进一步核查的城市空间。")

    add_section(doc, 2, "数据来源与研究对象")
    add_subsection(doc, "2.1", "北京 T-Drive 数据集")
    add_paragraph(doc, "T-Drive 数据集记录了北京出租车在一定时间范围内的 GPS 轨迹点，字段包括车辆编号、时间戳、经度和纬度。与订单型出租车数据不同，T-Drive 不直接提供乘客上车、下车、费用和订单状态，因此本文不能将单条记录解释为完整乘客订单，而是将其作为出租车运行轨迹观测点。")
    add_subsection(doc, "2.2", "研究对象界定")
    add_paragraph(doc, "本文研究对象定义为北京夜间出租车活动。所谓活动，是指根据轨迹点时间连续性、空间移动和停留状态提取出的轨迹活动段。活动段起点和终点分别用于近似表达夜间活动的出发与到达位置。为了避免概念误读，报告统一使用“活动起点”“活动终点”和“活动网格”，不将活动段直接解释为完整乘客订单。")
    add_subsection(doc, "2.3", "数据规模")
    add_table(
        doc,
        [
            ["指标", "数值"],
            ["输入轨迹记录数", fmt_int(summary.get("input_rows"))],
            ["清洗后记录数", fmt_int(summary.get("clean_rows"))],
            ["夜间活动记录数", fmt_int(summary.get("night_rows"))],
            ["参与分析活动网格数", fmt_int(summary.get("region_count"))],
            ["识别异常活动数", fmt_int(summary.get("anomaly_count"))],
        ],
    )

    add_section(doc, 3, "数据预处理与活动段构造")
    add_subsection(doc, "3.1", "坐标转换与空间过滤")
    add_paragraph(doc, "原始 T-Drive 坐标采用 WGS84 坐标系。由于前端地图使用高德地图，系统在数据处理过程中将轨迹点转换为 GCJ-02 坐标，并依据北京经纬度范围进行空间过滤，去除明显不在研究范围内的异常点。")
    add_subsection(doc, "3.2", "轨迹点到活动段")
    add_paragraph(doc, "系统按照车辆编号和时间戳对轨迹点排序，并根据相邻点时间间隔、速度跳变、长时间停留和活动段最大持续时长等条件切分轨迹。每个活动段包含起点、终点、持续时间和移动距离等信息。该步骤的目的不是还原真实订单，而是在缺少订单标签的条件下获得相对稳定的出租车活动分析单元。")
    add_image_placeholder(doc, "图1 数据预处理与活动段构造流程图（此处插入流程图片）")
    add_subsection(doc, "3.3", "夜间时段定义")
    add_paragraph(doc, "本文将夜间分析窗口设定为 20:00 至次日 03:00。该时段覆盖晚间出行、夜间消费、深夜返程和凌晨交通接驳等活动，能够较好对应城市夜间活力与夜间交通服务需求。")

    add_section(doc, 4, "H3 活动网格与地名增强")
    add_subsection(doc, "4.1", "H3 网格作为空间分析单元")
    add_paragraph(doc, "行政区尺度过大，难以表达夜间活动在街道、商圈和交通节点层面的差异。因此本文采用 H3 网格作为基础空间分析单元。H3 网格具有规则、可比较、便于聚合和适合地图渲染的特点，能够把出租车活动统一映射到稳定空间单元中。")
    add_subsection(doc, "4.2", "H3 分辨率选择")
    add_paragraph(doc, "项目实验过程中曾使用更高分辨率网格，但网格数量较多会导致前端渲染卡顿，并且容易覆盖底图信息。最终系统采用 H3 resolution 7，在空间表达精度、前端性能和地图可读性之间取得平衡。")
    add_subsection(doc, "4.3", "高德逆地理编码地名增强")
    add_paragraph(doc, "H3 编号本身缺乏可读性。为解决抽象网格难以解释的问题，系统对每个 H3 网格计算中心点，并调用高德逆地理编码接口，获取区县、街道、商圈、附近 POI 和 AOI 信息，形成 grid_name_lookup 命名表。最终前端优先显示“区县 · 街道”或“区县 · 商圈”等名称，使结果更符合北京城市空间认知。")
    add_image_placeholder(doc, "图2 H3 网格与高德逆地理编码地名增强示意图（此处插入地图或方法图片）")

    add_section(doc, 5, "夜间出租车活动特征体系")
    add_subsection(doc, "5.1", "活动规模特征")
    add_paragraph(doc, "活动规模特征包括活动起点数、活动终点数、活动总量和最大小时活动量，用于衡量某一网格在夜间出租车活动中的总体强度。")
    add_subsection(doc, "5.2", "时间结构特征")
    add_paragraph(doc, "时间结构特征包括深夜活动占比、活跃夜晚数、活跃小时数和周末增幅。这些指标用于区分稳定活跃区域、周末增强区域和凌晨仍保持活动的区域。")
    add_subsection(doc, "5.3", "空间流动特征")
    add_paragraph(doc, "空间流动特征包括活动起终点平衡、平均活动距离、平均持续时长和短活动比例，用于识别网格更偏出发、到达、短途接驳还是长距离交通联系。")
    add_subsection(doc, "5.4", "POI 与功能解释特征")
    add_paragraph(doc, "POI 与功能解释特征来自高德逆地理结果中的周边 POI、AOI 和商圈信息。系统按餐饮、购物、休闲、住宿、交通、科教文化、商务住宅、公司企业和公共服务等类型进行归类，并计算 POI 多样性、夜间相关 POI 数、交通设施 POI 数、商圈数量和 AOI 数量。")

    add_section(doc, 6, "夜间活力评分与区域类型识别")
    add_subsection(doc, "6.1", "Night Vitality Score 构建")
    add_paragraph(doc, "Night Vitality Score 综合活动规模、深夜占比、周末增幅、活动持续性和流动结构等指标。系统先对各项特征进行鲁棒归一化，再按照预设权重计算综合分数。该分数用于地图颜色渲染、网格排序、类型识别和异常解释。")
    add_subsection(doc, "6.2", "区域聚类与类型命名")
    add_paragraph(doc, "系统基于网格特征进行聚类，并结合指标表现为聚类结果赋予语义标签，包括夜间活动核心区、混合夜间活动区、夜间到达热点、夜间出发热点、交通枢纽型网格和低活跃区等。")
    add_subsection(doc, "6.3", "交通枢纽修正")
    add_paragraph(doc, "机场、车站和快速路节点可能因长距离交通需求产生高活动量，但其含义不同于夜间消费或城市活力核心。系统因此设置交通枢纽修正项，根据长距离活动比例、交通 POI 和区域名称线索识别交通枢纽型网格，降低将交通流量误读为夜生活活力的风险。")
    add_subsection(doc, "6.4", "评分方法对比")
    add_paragraph(doc, "除人工权重评分外，系统还生成熵权法和 PCA 分数，用于说明不同评分方法下结果的一致性和差异。该部分主要服务于可信度分析，而不是替代主评分体系。")
    add_table(doc, score_method_rows(scores))

    add_section(doc, 7, "高德 POI 功能画像")
    add_subsection(doc, "7.1", "POI 数据来源")
    add_paragraph(doc, "项目使用高德逆地理编码返回的周边 POI、AOI 和商圈信息作为网格功能画像来源。每个网格最多获取若干附近 POI，并记录其类型路径、距离和名称。")
    add_subsection(doc, "7.2", "POI 类型归类")
    add_paragraph(doc, "系统将高德 POI 中文类型归并为餐饮、购物、休闲、住宿、交通、文化、居住、企业和公共服务等功能组。相比旧的关键词推断方法，该方法直接利用高德返回的真实周边功能信息，能够更适配北京空间。")
    add_subsection(doc, "7.3", "POI 多样性指标")
    add_paragraph(doc, "POI 多样性综合考虑有效 POI 数量、功能类别数量、类别均衡度和覆盖度。重新计算后，全市活动网格 POI 多样性平均值约为 58.9%，中位数约为 67.5%，说明该指标已经能够区分不同网格的周边功能丰富程度。")
    add_table(
        doc,
        [
            ["指标", "数值"],
            ["POI 多样性均值", "58.9%"],
            ["POI 多样性中位数", "67.5%"],
            ["POI 多样性最大值", "98.5%"],
            ["POI 多样性为 0 的网格", "75 / 2357"],
        ],
    )

    add_section(doc, 8, "异常活动检测与 AI 辅助解释")
    add_subsection(doc, "8.1", "异常检测方法")
    add_paragraph(doc, "异常检测以同一网格同一小时的历史活动量为参照，使用 median 和 MAD 构造鲁棒基线，并计算异常强度 z-score。该方法能够减少极端值对基线的影响，更适合短时间窗口下的出租车活动异常识别。")
    add_subsection(doc, "8.2", "异常事件结果")
    add_paragraph(doc, f"系统共识别出 {fmt_int(summary.get('anomaly_count'))} 个异常夜间活动事件。异常事件主要表现为某一网格在特定夜间小时活动量明显高于历史同期基线。")
    add_table(doc, anomaly_rows(anomalies))
    add_subsection(doc, "8.3", "AI 辅助解释机制")
    add_paragraph(doc, "系统提供 AI 辅助解释功能，但解释严格限制在项目内数据范围内，包括异常活动量、历史基线、网格类型、POI 特征、活动起终点结构和小时上下文。AI 不直接编造外部新闻、天气或活动事件，只给出数据内可支持的可能原因和建议核查方向。")
    add_image_placeholder(doc, "图3 异常活动页面与 AI 分析结果截图（此处插入项目图片）")

    add_section(doc, 9, "系统设计与实现")
    add_subsection(doc, "9.1", "系统架构")
    add_paragraph(doc, "系统采用 Python 数据处理、Flask API、Vue 前端和高德地图 JS API 的三层架构。数据层负责轨迹清洗、活动段构造、H3 聚合、评分、POI 画像和异常检测；后端负责读取输出文件并提供接口；前端负责地图渲染、交互查询和分析结果展示。")
    add_image_placeholder(doc, "图4 系统总体架构图（此处插入架构图片）")
    add_subsection(doc, "9.2", "后端接口设计")
    add_paragraph(doc, "后端接口包括 summary、regions、geojson、hourly、anomalies、score-methods、forecast-metrics、forecasts、report 和 AI anomaly explain 等。前端通过这些接口获取统计摘要、网格列表、地图图层、小时曲线、异常事件、评分可信度和 AI 分析结果。")
    add_subsection(doc, "9.3", "前端功能设计")
    add_paragraph(doc, "前端包含活力地图、异常活动、网格画像和可信度四个主要视图。活力地图使用连续暖色渐变表达活动强度，只显示 Top 10 CORE 核心标记，并支持点击高亮、时间滑块、搜索定位和详情卡片。")
    add_image_placeholder(doc, "图5 北京夜间活力地图主界面截图（此处插入项目图片）")
    add_image_placeholder(doc, "图6 活动网格画像与可信度页面截图（此处插入项目图片）")

    add_section(doc, 10, "北京夜间出租车活动结果分析")
    add_subsection(doc, "10.1", "总体空间格局")
    add_paragraph(doc, "从结果看，北京夜间出租车活动并非均匀分布，而是在空港、商务区、高校周边、居住与交通混合区等空间形成较强活动集聚。高活力网格主要分布在朝阳、海淀、顺义、东城和丰台等区域，这些空间兼具交通联系、就业活动、居住返程或夜间消费等功能。")
    add_subsection(doc, "10.2", "Top 活动网格分析")
    add_table(doc, top_region_rows(scores, 10))
    add_subsection(doc, "10.3", "典型区域画像")
    add_table(doc, typical_region_rows(scores))
    add_paragraph(doc, "空港与交通活动型区域体现出较强的夜间接驳和长距离活动特征；商务活动型区域通常活动总量较高，POI 多样性也较强；高校与科研活动型区域表现出较稳定的晚间和深夜出租车需求；城市核心混合型区域则往往同时具有居住、办公、餐饮和交通功能。")
    add_subsection(doc, "10.4", "小时变化与周末效应")
    add_paragraph(doc, "20:00 至 03:00 的小时曲线能够反映晚间活动、深夜返程和凌晨接驳的节奏。部分网格周末增幅显著，说明其活动可能与休闲消费、聚会或非工作日出行有关；而交通枢纽或居住区周边则可能呈现更稳定的夜间活动结构。")
    add_image_placeholder(doc, "图7 Top 网格小时变化与周末效应截图（此处插入项目图片）")

    add_section(doc, 11, "可信度分析与局限性")
    add_subsection(doc, "11.1", "评分稳定性")
    add_paragraph(doc, "评分方法对比显示，不同方法对高活力区域的排序并不完全一致。人工权重强调可解释性和业务逻辑，熵权法强调数据差异性，PCA 强调整体方差方向。三者差异说明夜间活力并非单一维度指标，因此报告采用人工权重作为主解释指标，并将其他方法作为稳健性参考。")
    add_subsection(doc, "11.2", "预测模型表现")
    add_table(doc, forecast_rows())
    add_paragraph(doc, "预测模型部分用于评估短时活动量是否具有可学习规律。Random Forest 相比上一小时基线通常具有更低误差，但由于数据时间跨度较短，预测结果只作为可信度辅助分析，不作为项目主要结论。")
    add_subsection(doc, "11.3", "数据局限")
    add_paragraph(doc, "T-Drive 数据发布时间较早，时间跨度有限，并且不是订单数据，缺少乘客真实上车、下车、费用和目的信息。因此本文结论应被理解为基于出租车轨迹活动段的夜间活动监测，而非完整乘客出行行为研究。")
    add_subsection(doc, "11.4", "方法局限")
    add_paragraph(doc, "活动段切分依赖时间间隔、速度和停留阈值等规则，可能与真实运营状态存在偏差。高德逆地理返回的是网格中心附近 POI，不等同于完整 POI 普查。异常检测只能说明数据内偏离历史基线，不能直接确认真实事件原因。")

    add_section(doc, 12, "结论与展望")
    add_subsection(doc, "12.1", "研究结论")
    add_paragraph(doc, "本文完成了基于北京 T-Drive 轨迹数据的夜间出租车活动监测系统。系统能够从轨迹点构造活动段，将活动聚合到 H3 网格，通过高德逆地理编码增强地名可读性，并结合 POI 功能画像、活力评分、异常检测和可视化交互识别北京夜间出租车活动空间。")
    add_subsection(doc, "12.2", "系统价值")
    add_paragraph(doc, "该系统可为城市夜间交通供给、夜间活动空间识别、异常需求监测和城市空间治理提供辅助依据。相比静态行政区统计，H3 活动网格能够更细致地展示夜间活动差异；相比单纯热力图，系统进一步提供了地名解释、POI 画像、异常分析和可信度视图。")
    add_subsection(doc, "12.3", "后续工作")
    add_paragraph(doc, "后续可继续引入更多月份或更新年份的出租车、网约车和公共交通数据，补充天气、节假日和城市事件信息，接入更完整的 POI 数据源，并进一步优化前端地图性能、网格筛选和空间统计解释。若获得真实订单数据，还可以将活动段近似分析升级为乘客出行订单分析。")

    add_appendix_heading(doc)
    add_subsection(doc, "附录A", "系统运行命令")
    add_table(
        doc,
        [
            ["任务", "PowerShell 命令"],
            ["启动虚拟环境", r"cd D:\uns_2\final_2; .\.venv\Scripts\Activate.ps1"],
            ["启动后端", r"$env:PYTHONPATH='src'; $env:NIGHTSENSE_OUTPUT_DIR='outputs\beijing_tdrive'; python backend\app.py"],
            ["启动前端", r"cd D:\uns_2\final_2\frontend; npm.cmd run dev"],
            ["重新生成网格地名", r"python scripts\data\enrich_beijing_grid_names.py --sleep 0.18"],
            ["重新生成完整报告", r"python scripts\reports\generate_beijing_final_report.py"],
        ],
    )
    add_subsection(doc, "附录B", "主要输出文件")
    add_table(
        doc,
        [
            ["文件", "作用"],
            ["activity_segments_gcj02.csv", "由 T-Drive 轨迹点构造得到的活动段数据"],
            ["region_scores_phase2.csv", "融合 POI 与交通枢纽修正后的活动网格评分结果"],
            ["h3_region_scores_geojson.json", "前端地图使用的 H3 活动网格 GeoJSON"],
            ["grid_name_lookup.csv", "H3 网格中心点对应的高德逆地理命名表"],
            ["grid_name_lookup_raw.json", "高德逆地理编码原始响应缓存"],
            ["poi_features.csv", "基于高德 POI/AOI/商圈计算的功能画像指标"],
            ["anomalies_attributed.csv", "带可能原因说明的异常活动事件"],
            ["score_method_comparison.csv", "人工权重、熵权法与 PCA 评分对比"],
            ["forecast_metrics.csv", "短时活动量预测模型评价指标"],
        ],
    )
    add_subsection(doc, "附录C", "插图占位清单")
    add_table(
        doc,
        [
            ["图号", "建议插入内容"],
            ["图1", "数据预处理与活动段构造流程图"],
            ["图2", "H3 网格与高德逆地理编码地名增强示意图"],
            ["图3", "异常活动页面与 AI 分析结果截图"],
            ["图4", "系统总体架构图"],
            ["图5", "北京夜间活力地图主界面截图"],
            ["图6", "活动网格画像与可信度页面截图"],
            ["图7", "Top 网格小时变化与周末效应截图"],
        ],
    )

    add_paragraph(doc, first_line_indent=False)
    add_paragraph(
        doc,
        "参考资料与数据来源：北京 T-Drive 出租车轨迹数据；高德地图 Web JS API 与逆地理编码服务；Uber H3 空间索引；项目输出目录 outputs/beijing_tdrive。",
        first_line_indent=False,
    )

    force_all_text_black(doc)
    try:
        doc.save(REPORT_PATH)
        return REPORT_PATH
    except PermissionError:
        fallback = REPORT_PATH.with_name(f"{REPORT_PATH.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{REPORT_PATH.suffix}")
        doc.save(fallback)
        return fallback


if __name__ == "__main__":
    print(build_report())
