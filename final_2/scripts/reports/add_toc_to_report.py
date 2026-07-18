from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


REPORT = Path("reports/Urban_NightSense_阶段二项目报告.docx")


def set_style_font(doc: Document, style_name: str, size: int, bold: bool = True) -> None:
    style = doc.styles[style_name]
    style.font.name = "宋体"
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def make_text_paragraph(text: str, size_half_points: str = "24", bold: bool = False, center: bool = False):
    paragraph = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    if center:
        just = OxmlElement("w:jc")
        just.set(qn("w:val"), "center")
        ppr.append(just)
    paragraph.append(ppr)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if bold:
        rpr.append(OxmlElement("w:b"))
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), size_half_points)
    rpr.append(size)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "宋体")
    rpr.append(fonts)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    return paragraph


def make_toc_field():
    paragraph = OxmlElement("w:p")
    for tag, text in [
        ("begin", None),
        ("instr", 'TOC \\o "1-2" \\h \\z \\u'),
        ("separate", None),
        ("placeholder", "请在 Word 中右键选择“更新域/更新目录”，生成目录页码。"),
        ("end", None),
    ]:
        run = OxmlElement("w:r")
        if tag in {"begin", "separate", "end"}:
            field = OxmlElement("w:fldChar")
            field.set(qn("w:fldCharType"), tag)
            run.append(field)
        elif tag == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run.append(instr)
        else:
            text_node = OxmlElement("w:t")
            text_node.text = text
            run.append(text_node)
        paragraph.append(run)
    return paragraph


def make_page_break():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def remove_existing_generated_toc(doc: Document) -> None:
    body = doc._body._element
    for p in list(doc.paragraphs):
        if p.text.strip() != "目录":
            continue
        element = p._element
        following = []
        node = element.getnext()
        while node is not None:
            following.append(node)
            if node.tag == qn("w:p") and re.match(r"^[一二三四五六七八九十]+、", paragraph_text(node)):
                break
            node = node.getnext()
        body.remove(element)
        for item in following[:-1]:
            if item.getparent() is body:
                body.remove(item)
        return


def apply_heading_styles(doc: Document) -> None:
    set_style_font(doc, "Heading 1", 14, True)
    set_style_font(doc, "Heading 2", 12, True)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^[一二三四五六七八九十]+、", text):
            paragraph.style = doc.styles["Heading 1"]
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(8)
        elif re.match(r"^\d+\.\d+\s+", text):
            paragraph.style = doc.styles["Heading 2"]
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(6)


def add_toc(doc: Document) -> None:
    body = doc._body._element
    first_heading = None
    for paragraph in doc.paragraphs:
        if re.match(r"^[一二三四五六七八九十]+、", paragraph.text.strip()):
            first_heading = paragraph._element
            break
    if first_heading is None:
        raise RuntimeError("未找到正文一级标题，无法插入目录。")

    toc_title = make_text_paragraph("目录", size_half_points="28", bold=True, center=True)
    toc_field = make_toc_field()
    page_break = make_page_break()

    body.insert(body.index(first_heading), page_break)
    body.insert(body.index(page_break), toc_field)
    body.insert(body.index(toc_field), toc_title)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    if not REPORT.exists():
        raise FileNotFoundError(REPORT)

    backup = REPORT.with_name(f"{REPORT.stem}_before_toc_{datetime.now().strftime('%Y%m%d_%H%M%S')}{REPORT.suffix}")
    shutil.copy2(REPORT, backup)

    doc = Document(REPORT)
    remove_existing_generated_toc(doc)
    apply_heading_styles(doc)
    add_toc(doc)
    enable_update_fields(doc)
    doc.save(REPORT)
    print(f"updated={REPORT}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
